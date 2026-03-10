# data_loader_management/data_loader_manager.py
# -*- coding: utf-8 -*-
"""
DataLoaderManager — SOLID uyumlu, “yalın” veri yükleyici.

Amaç:
- Eğitim (QA) modunda: JSON dosyalarındaki Soru–Cevap çiftlerini ham hâliyle döndürür.
- İnference modunda: Düz metin (txt/json içi string listesi) ham hâliyle döndürür.
- Hiçbir normalize / lowercase / tag / tokenize / tensorize YAPMAZ.
- Fail-fast (strict=True) yaklaşımıyla sessiz hataları ortadan kaldırır.

Sözleşmeler:
- QA_TRAIN modunda yalnızca .json dosyaları kabul edilir.
  Beklenen kök: List[Dict]. Varsayılan anahtarlar: "Soru"/"Cevap".
  Alternatif anahtarlar: "question"/"answer" (parametrelerle genişletilebilir).

- TEXT_INFER modunda .txt ve .json kabul edilir.
  .txt -> satırların tamamı tek bir string olarak okunur (ham).
  .json -> kök List[str] ise her eleman bir metin girdisidir.

Çıkışlar:
- load() → mode=QA_TRAIN: List[Tuple[str, str]]  # (question, answer)
- load() → mode=TEXT_INFER: List[str]            # metinler

Not:
- İleride başka modaliteler eklenecekse Loader arayüzü/registry genişletilebilir.
- Bu modül TokenizerCore’a bağımlı değildir; tam tersi de geçerli.
"""

from __future__ import annotations
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Tuple, Dict, Any, Iterable, Optional
import json
import logging
import re
import hashlib

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

logger = logging.getLogger(__name__)
if not logger.handlers:
    _h = logging.StreamHandler()
    _h.setFormatter(logging.Formatter("[%(levelname)s][DataLoaderManager] %(message)s"))
    logger.addHandler(_h)
logger.setLevel(logging.INFO)


# ===========================
# Hatalar
# ===========================
class DataLoaderError(Exception): ...
class DataDirectoryNotFoundError(DataLoaderError): ...
class UnsupportedFormatError(DataLoaderError): ...
class InvalidStructureError(DataLoaderError): ...
class EmptyTextError(DataLoaderError): ...


# ===========================
# Mod & Konfig
# ===========================
class LoadMode:
    QA_TRAIN = "QA_TRAIN"      # JSON dosyaları → soru-cevap çiftleri
    TEXT_INFER = "TEXT_INFER"  # TXT/JSON dosyaları → raw text
    RAW_TEXT = "RAW_TEXT"      # TXT/DOCX dosyaları → raw text chunks


@dataclass(frozen=True)
class DataLoaderConfig:
    data_dir: Path
    mode: str = LoadMode.QA_TRAIN
    strict: bool = False                    # fail-fast
    # QA için anahtar eşleşmeleri (öncelik sırasıyla)
    question_keys: Iterable[str] = field(default_factory=lambda: ("Soru", "question", "instruction", "prompt"))
    answer_keys: Iterable[str]   = field(default_factory=lambda: ("Cevap", "answer", "output", "response"))
    # Dosya filtreleri (moda göre uygulanır)
    qa_extensions: Iterable[str]   = field(default_factory=lambda: (".json",))  # SADECE JSON!
    infer_extensions: Iterable[str]= field(default_factory=lambda: (".txt", ".docx"))  # JSON QA_TRAIN mode'da zaten yükleniyor
    raw_text_extensions: Iterable[str] = field(default_factory=lambda: (".txt", ".docx"))
    # Yürüyüş davranışı
    follow_symlinks: bool = False
    max_files: Optional[int] = None         # None: sınırsız
    # Akıllı bölme ayarları
    max_tokens_per_text: int = 198          #  DÜZELTME: Token sınırı (400'den 300'e düşürüldü - token tahmini yanlış olduğu için güvenli)
    enable_smart_splitting: bool = True     # Akıllı bölme aktif
    overlap_tokens: int = 20                # Parçalar arası örtüşme


# ===========================
# Ana Yönetici
# ===========================
class DataLoaderManager:
    def __init__(self, cfg: DataLoaderConfig):
        self.cfg = cfg
        self._stats = {
            "files_seen": 0,
            "files_loaded": 0,
            "examples": 0,
        }
        self._file_index_counter = 0  # ✅ File index counter (her dosya için unique ID)

    @property
    def stats(self) -> Dict[str, int]:
        return dict(self._stats)
    
    # ---- Akıllı Bölme Metodları -----------------------------------------------
    def _estimate_token_count(self, text: str) -> int:
        """
        Metnin yaklaşık token sayısını tahmin eder.
        Türkçe için daha konservatif tahmin (BPE tokenization için)
        """
        if not text or not text.strip():
            return 0
        
        #  KRİTİK DÜZELTME: Çok daha konservatif tahmin (BPE tokenization için)
        # BPE, kelimeleri çok daha fazla token'a bölebilir (özellikle uzun kelimeler, OOV kelimeler)
        # Örnek: "muhammed" → ["muham", "med</w>"] (2 token)
        # Örnek: "yapayzeka" → ["yap", "ay", "zek", "a</w>"] (4 token)
        # Gerçek durum: 400 token tahmin → 1278 token gerçek (3.2x fark!)
        # Bu yüzden çok daha konservatif tahmin: kelime sayısı * 2.0 (güvenli)
        words = len(text.split())
        estimated_tokens = int(words * 2.5)  #  1.5 → 2.3 (çok daha konservatif - gerçek token sayısına yakın)
        return estimated_tokens
    
    def _smart_split_text(self, text: str, max_tokens: int = 198, overlap: int = 20) -> List[str]:
        """
        Uzun metni akıllıca böler.
        - Cümle sınırlarında böler
        - Parçalar arası örtüşme sağlar
        - Hiçbir veri kaybı olmaz
        - Duplicate chunk'ları filtreler (kaynak verideki tekrar eden içerik için)
        """
        if not text or not text.strip():
            return []
        
        estimated_tokens = self._estimate_token_count(text)
        
        # Eğer metin zaten kısa ise, bölme
        if estimated_tokens <= max_tokens:
            return [text.strip()]
        
        # Cümleleri ayır
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        if not sentences:
            return [text.strip()]
        
        chunks = []
        current_chunk = ""
        current_tokens = 0
        
        for sentence in sentences:
            sentence_tokens = self._estimate_token_count(sentence)
            
            # Eğer tek cümle bile sınırı aşıyorsa, kelime bazında böl
            if sentence_tokens > max_tokens:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                    current_chunk = ""
                    current_tokens = 0
                
                # Cümleyi kelime bazında böl
                word_chunks = self._split_by_words(sentence, max_tokens, overlap)
                chunks.extend(word_chunks)
                continue
            
            # Normal akış: cümle ekle
            if current_tokens + sentence_tokens > max_tokens:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                
                # Örtüşme için son kelimeleri al
                overlap_text = self._get_overlap_text(current_chunk, overlap)
                current_chunk = overlap_text + " " + sentence
                current_tokens = self._estimate_token_count(current_chunk)
            else:
                if current_chunk:
                    current_chunk += ". " + sentence
                else:
                    current_chunk = sentence
                current_tokens += sentence_tokens
        
        # Son parçayı ekle
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        # ✅ DEDUPLICATION: Kaynak verideki tekrar eden içeriği filtrele
        # Overlap'i hesaba katarak, sadece tam duplicate chunk'ları filtrele
        filtered_chunks = self._deduplicate_chunks(chunks, overlap)
        
        return filtered_chunks
    
    def _deduplicate_chunks(self, chunks: List[str], overlap: int) -> List[str]:
        """
        Chunk listesinden duplicate'leri filtrele.
        Overlap'i hesaba katar - overlap nedeniyle benzer chunk'lar normal kabul edilir.
        Sadece tam duplicate chunk'lar (overlap dışında da aynı) filtrelenir.
        """
        if not chunks:
            return []
        
        seen_hashes = set()
        filtered_chunks = []
        duplicate_count = 0
        
        for chunk in chunks:
            chunk_stripped = chunk.strip()
            if not chunk_stripped:
                continue
            
            # Chunk'ı hash'le (tam içerik)
            chunk_hash = hashlib.sha256(chunk_stripped.encode('utf-8')).hexdigest()
            
            # Eğer bu hash daha önce görüldüyse, duplicate
            if chunk_hash in seen_hashes:
                duplicate_count += 1
                continue  # Duplicate chunk'ı atla
            
            # Normal chunk - ekle
            seen_hashes.add(chunk_hash)
            filtered_chunks.append(chunk_stripped)
        
        if duplicate_count > 0:
            # Logger varsa kullan, yoksa print
            try:
                import logging
                logger = logging.getLogger(__name__)
                logger.debug(f"[DataLoaderManager] {duplicate_count} duplicate chunk filtrelendi (toplam {len(chunks)} chunk)")
            except:
                pass  # Logger yoksa sessizce devam et
        
        return filtered_chunks
    
    def _split_by_words(self, text: str, max_tokens: int, overlap: int) -> List[str]:
        """Metni kelime bazında böler"""
        words = text.split()
        if not words:
            return []
        
        chunks = []
        current_chunk = []
        current_tokens = 0
        
        for word in words:
            word_tokens = self._estimate_token_count(word)
            
            if current_tokens + word_tokens > max_tokens:
                if current_chunk:
                    chunks.append(" ".join(current_chunk))
                
                # Örtüşme için son kelimeleri al
                overlap_words = current_chunk[-overlap:] if len(current_chunk) >= overlap else current_chunk
                current_chunk = overlap_words + [word]
                current_tokens = self._estimate_token_count(" ".join(current_chunk))
            else:
                current_chunk.append(word)
                current_tokens += word_tokens
        
        if current_chunk:
            chunks.append(" ".join(current_chunk))
        
        return chunks
    
    def _get_overlap_text(self, text: str, overlap_tokens: int) -> str:
        """Metnin sonundan belirtilen token sayısı kadar alır"""
        if not text:
            return ""
        
        words = text.split()
        if len(words) <= overlap_tokens:
            return text
        
        # Son N kelimeyi al
        overlap_words = words[-overlap_tokens:]
        return " ".join(overlap_words)
    
    def _process_qa_pair_with_splitting(self, question: str, answer: str) -> List[Tuple[str, str]]:
        """
        Soru-Cevap çiftini akıllı bölme ile işler.
        Uzun cevapları böler, soruları korur.
        """
        if not self.cfg.enable_smart_splitting:
            return [(question, answer)]
        
        max_tokens = self.cfg.max_tokens_per_text
        overlap = self.cfg.overlap_tokens
        
        # Soru uzunluğunu kontrol et
        question_tokens = self._estimate_token_count(question)
        answer_tokens = self._estimate_token_count(answer)
        
        # Eğer soru çok uzunsa, onu da böl
        if question_tokens > max_tokens:
            question_chunks = self._smart_split_text(question, max_tokens, overlap)
            if len(question_chunks) > 1:
                # Uzun soruyu kısalt
                question = question_chunks[0] + "..."
        
        # Cevap uzunluğunu kontrol et
        if answer_tokens <= max_tokens:
            return [(question, answer)]
        
        # Cevabı böl
        answer_chunks = self._smart_split_text(answer, max_tokens, overlap)
        
        # Her parça için ayrı QA çifti oluştur
        qa_pairs = []
        for i, chunk in enumerate(answer_chunks):
            if i == 0:
                # İlk parça: orijinal soru
                qa_pairs.append((question, chunk))
            else:
                # Sonraki parçalar: devam soruları
                continuation_question = f"{question} (devam {i+1})"
                qa_pairs.append((continuation_question, chunk))
        
        return qa_pairs

    # ---- Public API ----------------------------------------------------------
    def load(self) -> List[Tuple[str, str]] | List[str]:
        """
        Moda göre uygun yükleyiciyi çağırır.
        QA_TRAIN  → List[Tuple[question, answer]]
        TEXT_INFER→ List[str]
        """
        if not self.cfg.data_dir.exists():
            raise DataDirectoryNotFoundError(f"Veri klasörü yok: {self.cfg.data_dir}")

        if self.cfg.mode == LoadMode.QA_TRAIN:
            return self._load_qa_pairs()
        elif self.cfg.mode == LoadMode.TEXT_INFER:
            return self._load_text_inputs()
        elif self.cfg.mode == LoadMode.RAW_TEXT:
            return self._load_raw_text_chunks()
        else:
            raise DataLoaderError(f"Bilinmeyen mode: {self.cfg.mode}")
    
    def load_with_file_index(self) -> List[Tuple[str, str, int]] | List[Tuple[str, int]]:
        """
        ✅ YENİ: Dosya indeksi ile yükle (overlap önleme için)
        
        Returns:
            - QA_TRAIN: List[Tuple[question, answer, file_idx]]
            - RAW_TEXT: List[Tuple[chunk, file_idx]]
            - TEXT_INFER: List[Tuple[text, file_idx]]
        """
        if not self.cfg.data_dir.exists():
            raise DataDirectoryNotFoundError(f"Veri klasörü yok: {self.cfg.data_dir}")

        if self.cfg.mode == LoadMode.QA_TRAIN:
            return self._load_qa_pairs_with_file_index()
        elif self.cfg.mode == LoadMode.TEXT_INFER:
            return self._load_text_inputs_with_file_index()
        elif self.cfg.mode == LoadMode.RAW_TEXT:
            return self._load_raw_text_chunks_with_file_index()
        else:
            raise DataLoaderError(f"Bilinmeyen mode: {self.cfg.mode}")

    # ---- Internal: QA (Training) --------------------------------------------
    def _load_qa_pairs(self) -> List[Tuple[str, str]]:
        out: List[Tuple[str, str]] = []
        files = self._iter_files(allowed_exts=set(map(str.lower, self.cfg.qa_extensions)))

        for file_path in files:
            loaded_from_this_file = 0
            
            if file_path.suffix.lower() == '.json':
                data = self._read_json(file_path)
                if not isinstance(data, list):
                    self._handle_error(InvalidStructureError(f"{file_path.name} kök yapı List olmalı; {type(data)}"))
                    continue

                for idx, row in enumerate(data):
                    if not isinstance(row, dict):
                        self._handle_error(InvalidStructureError(f"{file_path.name}[{idx}] öğesi dict olmalı; {type(row)}"))
                        continue

                    # Özel format kontrolü
                    if self._is_sentiment_format(row):
                        q, a = self._convert_sentiment_to_qa(row)
                    elif self._is_instruction_format(row):
                        q, a = self._convert_instruction_to_qa(row)
                    else:
                        # Standart QA formatı
                        q_key = self._resolve_key(row, self.cfg.question_keys, kind="Soru", file=file_path.name, idx=idx)
                        a_key = self._resolve_key(row, self.cfg.answer_keys,   kind="Cevap", file=file_path.name, idx=idx)

                        q_raw = row.get(q_key)
                        a_raw = row.get(a_key)

                        if not isinstance(q_raw, str) or not isinstance(a_raw, str):
                            self._handle_error(InvalidStructureError(f"{file_path.name}[{idx}] Soru/Cevap string olmalı."))
                            continue

                        q = q_raw
                        a = a_raw

                    if len(q.strip()) == 0 or len(a.strip()) == 0:
                        self._handle_error(EmptyTextError(f"{file_path.name}[{idx}] boş (yalnızca boşluk) Soru/Cevap."))
                        continue

                    # Akıllı bölme ile işle
                    processed_pairs = self._process_qa_pair_with_splitting(q, a)
                    for processed_q, processed_a in processed_pairs:
                        out.append((processed_q, processed_a))
                        loaded_from_this_file += 1
                        self._stats["examples"] += 1

            logger.info(f"[QA] {file_path.name} → {loaded_from_this_file} çift")

        logger.info(f"[QA DONE] files_loaded={self._stats['files_loaded']} examples={self._stats['examples']}")
        return out
    
    def _load_qa_pairs_with_file_index(self) -> List[Tuple[str, str, int]]:
        """
        ✅ YENİ: QA çiftlerini dosya indeksi ile yükle (overlap önleme için)
        
        Returns:
            List[Tuple[str, str, int]]: (question, answer, file_index)
        """
        out: List[Tuple[str, str, int]] = []
        
        # ✅ DÜZELTME: _iter_files_with_index kullan (unique, tutarlı file_idx için)
        files_with_idx = self._iter_files_with_index(allowed_exts=set(map(str.lower, self.cfg.qa_extensions)))
        
        file_count = 0
        for file_path, file_idx in files_with_idx:  # ✅ file_idx tracking (tutarlı)
            file_count += 1
            loaded_from_this_file = 0
            
            if file_path.suffix.lower() == '.json':
                data = self._read_json(file_path)
                if not isinstance(data, list):
                    self._handle_error(InvalidStructureError(f"{file_path.name} kök yapı List olmalı; {type(data)}"))
                    continue

                for idx, row in enumerate(data):
                    if not isinstance(row, dict):
                        self._handle_error(InvalidStructureError(f"{file_path.name}[{idx}] öğesi dict olmalı; {type(row)}"))
                        continue

                    # Özel format kontrolü
                    if self._is_sentiment_format(row):
                        q, a = self._convert_sentiment_to_qa(row)
                    elif self._is_instruction_format(row):
                        q, a = self._convert_instruction_to_qa(row)
                    else:
                        # Standart QA formatı
                        q_key = self._resolve_key(row, self.cfg.question_keys, kind="Soru", file=file_path.name, idx=idx)
                        a_key = self._resolve_key(row, self.cfg.answer_keys, kind="Cevap", file=file_path.name, idx=idx)

                        q_raw = row.get(q_key)
                        a_raw = row.get(a_key)

                        if not isinstance(q_raw, str) or not isinstance(a_raw, str):
                            self._handle_error(InvalidStructureError(f"{file_path.name}[{idx}] Soru/Cevap string olmalı."))
                            continue

                        q = q_raw
                        a = a_raw

                    if len(q.strip()) == 0 or len(a.strip()) == 0:
                        self._handle_error(EmptyTextError(f"{file_path.name}[{idx}] boş (yalnızca boşluk) Soru/Cevap."))
                        continue

                    # Akıllı bölme ile işle
                    processed_pairs = self._process_qa_pair_with_splitting(q, a)
                    
                    # ✅ Her QA çiftine file_idx ekle
                    for processed_q, processed_a in processed_pairs:
                        out.append((processed_q, processed_a, file_idx))  # ✅ file_idx ekle
                        loaded_from_this_file += 1
                        self._stats["examples"] += 1

            logger.info(f"[QA+FileIdx] [{file_idx}] {file_path.name} → {loaded_from_this_file} çift")

        logger.info(
            f"[QA+FileIdx DONE] files_loaded={file_count}, "
            f"examples={len(out)}, "
            f"avg_qa_per_file={len(out)/file_count if file_count > 0 else 0:.1f}"
        )
        return out
    
    def _load_text_inputs_with_file_index(self) -> List[Tuple[str, int]]:
        """
        ✅ YENİ: Text inputs'ları dosya indeksi ile yükle
        
        Returns:
            List[Tuple[str, int]]: (text, file_index)
        """
        out: List[Tuple[str, int]] = []
        
        # ✅ DÜZELTME: _iter_files_with_index kullan (unique, tutarlı file_idx için)
        files_with_idx = self._iter_files_with_index(allowed_exts=set(map(str.lower, self.cfg.infer_extensions)))
        
        for p, file_idx in files_with_idx:  # ✅ file_idx tracking (tutarlı)
            if p.suffix.lower() == ".txt":
                text = p.read_text(encoding="utf-8", errors="strict").strip()
                if len(text) > 0:
                    out.append((text, file_idx))
                    self._stats["examples"] += 1
            elif p.suffix.lower() == ".json":
                data = self._read_json(p)
                if isinstance(data, list):
                    for item in data:
                        if isinstance(item, str) and len(item.strip()) > 0:
                            out.append((item.strip(), file_idx))
                            self._stats["examples"] += 1
        
        return out

    # ---- Internal: Text (Inference) -----------------------------------------
    def _load_text_inputs(self) -> List[str]:
        out: List[str] = []
        files = self._iter_files(allowed_exts=set(map(str.lower, self.cfg.infer_extensions)))

        for p in files:
            if p.suffix.lower() == ".txt":
                text = p.read_text(encoding="utf-8", errors="strict")
                if len(text) == 0:
                    self._handle_error(EmptyTextError(f"{p.name} boş içerik."))
                    continue
                
                # Raw text'i chunklara böl (hibrit eğitim için)
                chunks = self._smart_split_text(text, self.cfg.max_tokens_per_text, self.cfg.overlap_tokens)
                out.extend(chunks)
                self._stats["examples"] += len(chunks)
                logger.info(f"[INF] {p.name} → {len(chunks)} chunk")

            elif p.suffix.lower() == ".docx":
                # DOCX dosyasını raw text olarak işle
                raw_text = self._read_docx_raw(p)
                if raw_text.strip():
                    # Raw text'i chunklara böl
                    chunks = self._smart_split_text(raw_text, self.cfg.max_tokens_per_text, self.cfg.overlap_tokens)
                    out.extend(chunks)
                    self._stats["examples"] += len(chunks)
                    logger.info(f"[INF] {p.name} → {len(chunks)} chunk")
            else:
                self._handle_error(UnsupportedFormatError(f"{p.name} desteklenmeyen uzantı (inference)."))

        logger.info(f"[INF DONE] files_loaded={self._stats['files_loaded']} examples={self._stats['examples']}")
        return out

    # ---- Internal: Raw Text Chunks -------------------------------------------
    def _load_raw_text_chunks(self) -> List[str]:
        """TXT/DOCX dosyalarından RAW text chunk'ları yükle"""
        out: List[str] = []
        files = self._iter_files(allowed_exts=set(map(str.lower, self.cfg.raw_text_extensions)))

        for file_path in files:
            loaded_from_this_file = 0
            
            if file_path.suffix.lower() == '.txt':
                # TXT dosyasını raw text chunk olarak işle
                try:
                    content = file_path.read_text(encoding="utf-8", errors="ignore").strip()
                    if content:
                        # Raw text'i chunklara böl
                        chunks = self._smart_split_text(content, self.cfg.max_tokens_per_text, self.cfg.overlap_tokens)
                        out.extend(chunks)
                        loaded_from_this_file = len(chunks)
                        self._stats["examples"] += len(chunks)
                except Exception as e:
                    self._handle_error(InvalidStructureError(f"{file_path.name} TXT okuma hatası: {e}"))
                    
            elif file_path.suffix.lower() == '.docx':
                # DOCX dosyasını raw text chunk olarak işle
                raw_text = self._read_docx_raw(file_path)
                if raw_text.strip():
                    # Raw text'i chunklara böl
                    chunks = self._smart_split_text(raw_text, self.cfg.max_tokens_per_text, self.cfg.overlap_tokens)
                    out.extend(chunks)
                    loaded_from_this_file = len(chunks)
                    self._stats["examples"] += len(chunks)
            
            logger.info(f"[RAW] {file_path.name} → {loaded_from_this_file} chunk")

        logger.info(f"[RAW DONE] files_loaded={self._stats['files_loaded']} examples={self._stats['examples']}")
        return out
    
    def _load_raw_text_chunks_with_file_index(self) -> List[Tuple[str, int]]:
        """
        ✅ YENİ: TXT/DOCX dosyalarından RAW text chunk'ları dosya indeksi ile yükle
        Her chunk için hangi dosyadan geldiğini track eder (overlap önleme için)
        
        Returns:
            List[Tuple[str, int]]: (chunk_text, file_index)
        """
        out: List[Tuple[str, int]] = []
        
        # ✅ DÜZELTME: _iter_files_with_index kullan (unique, tutarlı file_idx için)
        files_with_idx = self._iter_files_with_index(allowed_exts=set(map(str.lower, self.cfg.raw_text_extensions)))
        
        file_count = 0
        for file_path, file_idx in files_with_idx:  # ✅ file_idx tracking (tutarlı)
            file_count += 1
            loaded_from_this_file = 0
            
            if file_path.suffix.lower() == '.txt':
                # TXT dosyasını raw text chunk olarak işle
                try:
                    content = file_path.read_text(encoding="utf-8", errors="ignore").strip()
                    if content:
                        # Raw text'i chunklara böl
                        chunks = self._smart_split_text(content, self.cfg.max_tokens_per_text, self.cfg.overlap_tokens)
                        
                        # ✅ Her chunk'a file_idx ekle
                        for chunk in chunks:
                            out.append((chunk, file_idx))
                            loaded_from_this_file += 1
                        
                        self._stats["examples"] += len(chunks)
                except Exception as e:
                    self._handle_error(InvalidStructureError(f"{file_path.name} TXT okuma hatası: {e}"))
                    
            elif file_path.suffix.lower() == '.docx':
                # DOCX dosyasını raw text chunk olarak işle
                raw_text = self._read_docx_raw(file_path)
                if raw_text.strip():
                    # Raw text'i chunklara böl
                    chunks = self._smart_split_text(raw_text, self.cfg.max_tokens_per_text, self.cfg.overlap_tokens)
                    
                    # ✅ Her chunk'a file_idx ekle
                    for chunk in chunks:
                        out.append((chunk, file_idx))
                        loaded_from_this_file += 1
                    
                    self._stats["examples"] += len(chunks)
            
            logger.info(f"[RAW+FileIdx] [{file_idx}] {file_path.name} → {loaded_from_this_file} chunk")

        logger.info(
            f"[RAW+FileIdx DONE] files_loaded={file_count}, "
            f"examples={len(out)}, "
            f"avg_chunks_per_file={len(out)/file_count if file_count > 0 else 0:.1f}"
        )
        return out

    # ---- FS helpers ----------------------------------------------------------
    def _iter_files(self, allowed_exts: set[str]) -> Iterable[Path]:
        count = 0
        for p in sorted(self.cfg.data_dir.rglob("*") if self.cfg.follow_symlinks else self.cfg.data_dir.glob("**/*")):
            if not p.is_file():
                continue
            ext = p.suffix.lower()
            self._stats["files_seen"] += 1
            if ext not in allowed_exts:
                msg = f"Uzantı atlandı: {p.name} ({ext})"
                if self.cfg.strict:
                    raise UnsupportedFormatError(msg)
                logger.debug(msg)
                continue
            self._stats["files_loaded"] += 1
            yield p
            count += 1
            if self.cfg.max_files is not None and count >= self.cfg.max_files:
                logger.info(f"Maksimum dosya sayısına ulaşıldı: {self.cfg.max_files}")
                break
    
    def _iter_files_with_index(self, allowed_exts: set[str]) -> Iterable[Tuple[Path, int]]:
        """
        ✅ YENİ: Dosyaları (Path, file_index) tuple'ları olarak döndür
        Her dosya için unique, tutarlı file_index atar (overlap önleme için)
        """
        for p in sorted(self.cfg.data_dir.rglob("*") if self.cfg.follow_symlinks else self.cfg.data_dir.glob("**/*")):
            if not p.is_file():
                continue
            ext = p.suffix.lower()
            self._stats["files_seen"] += 1
            if ext not in allowed_exts:
                msg = f"Uzantı atlandı: {p.name} ({ext})"
                if self.cfg.strict:
                    raise UnsupportedFormatError(msg)
                logger.debug(msg)
                continue
            self._stats["files_loaded"] += 1
            file_idx = self._file_index_counter
            self._file_index_counter += 1
            yield (p, file_idx)
            if self.cfg.max_files is not None and self._file_index_counter >= self.cfg.max_files:
                logger.info(f"Maksimum dosya sayısına ulaşıldı: {self.cfg.max_files}")
                break

    def _read_json(self, path: Path) -> Any:
        try:
            return json.loads(path.read_text(encoding="utf-8", errors="strict"))
        except Exception as e:
            self._handle_error(InvalidStructureError(f"{path.name} JSON parse hatası: {e}"))
            return []  # strict=False ise akış devam edebilsin

    def _read_docx_raw(self, path: Path) -> str:
        """DOCX dosyasından sadece raw text çıkarır."""
        if not DOCX_AVAILABLE:
            self._handle_error(UnsupportedFormatError("python-docx kütüphanesi kurulu değil"))
            return ""
        
        try:
            doc = Document(str(path))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:  # Boş paragrafları atla
                    text_parts.append(text)
            
            # Paragrafları birleştir
            full_text = ' '.join(text_parts)
            
            logger.info(f"📄 {path.name}: {len(full_text)} karakter, ~{len(full_text.split()) * 1.3:.0f} token")
            return full_text
            
        except Exception as e:
            self._handle_error(InvalidStructureError(f"{path.name} DOCX okuma hatası: {e}"))
            return ""

    def _read_docx(self, path: Path) -> List[Dict[str, str]]:
        """DOCX dosyasını okuyup Soru-Cevap çiftlerine dönüştürür."""
        if not DOCX_AVAILABLE:
            self._handle_error(UnsupportedFormatError("python-docx kütüphanesi kurulu değil"))
            return []
        
        try:
            # Basit DOCX işleme
            return self._simple_docx_processing(path)
            
        except Exception as e:
            self._handle_error(InvalidStructureError(f"{path.name} DOCX okuma hatası: {e}"))
            return []
    
    def _simple_docx_processing(self, path: Path) -> List[Dict[str, str]]:
        """Basit DOCX işleme (fallback) - Autoregressive format için düzeltildi"""
        try:
            doc = Document(str(path))
            qa_pairs = []
            
            # Basit yaklaşım: Her paragrafı ayrı bir metin olarak işle
            current_text = ""
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    current_text += text + " "
            
            # Eğer metin varsa, autoregressive format için düz metin olarak işle
            if current_text.strip():
                # DOCX'teki düz metinleri autoregressive format için kullan
                # Soru: "Bu metin hakkında ne düşünüyorsun?"
                # Cevap: DOCX'teki metin
                question = "Bu metin hakkında ne düşünüyorsun?"
                answer = current_text.strip()
                if answer:
                        qa_pairs.append({
                            "Soru": question,
                            "Cevap": answer
                        })
                else:
                    # Tek cümle varsa, genel bir soru-cevap çifti oluştur
                    qa_pairs.append({
                        "Soru": "Bu metin hakkında ne öğrenebilirim?",
                        "Cevap": current_text.strip()
                    })
            
            return qa_pairs
            
        except Exception as e:
            self._handle_error(InvalidStructureError(f"{path.name} basit DOCX okuma hatası: {e}"))
            return []

    def _resolve_key(self, row: Dict[str, Any], candidates: Iterable[str], *, kind: str, file: str, idx: int) -> str:
        for k in candidates:
            if k in row:
                return k
        self._handle_error(InvalidStructureError(
            f"{file}[{idx}] {kind} anahtarı bulunamadı. Adaylar: {list(candidates)}"
        ))
        # strict=False ise akışa devam edebilmek için “yok” döndürmek yerine exception yükseltildi.
        # Bu noktaya akış düşmez çünkü _handle_error(strict=True) raise eder.
        raise InvalidStructureError(f"{file}[{idx}] {kind} anahtarı eksik.")

    def _is_sentiment_format(self, row: Dict[str, Any]) -> bool:
        """Sentiment analizi formatını kontrol eder"""
        return "text" in row and ("label" in row or "sentiment" in row)
    
    def _is_instruction_format(self, row: Dict[str, Any]) -> bool:
        """Instruction formatını kontrol eder"""
        return "instruction" in row and "output" in row
    
    def _convert_sentiment_to_qa(self, row: Dict[str, Any]) -> Tuple[str, str]:
        """Sentiment verisini QA formatına çevirir"""
        text = row.get("text", "")
        sentiment = row.get("sentiment", row.get("label", "bilinmeyen"))
        
        question = f"Bu metnin duygu analizi nasıldır?"
        answer = f"Metin: '{text}' - Duygu: {sentiment}"
        
        return question, answer
    
    def _convert_instruction_to_qa(self, row: Dict[str, Any]) -> Tuple[str, str]:
        """Instruction verisini QA formatına çevirir"""
        instruction = row.get("instruction", "")
        output = row.get("output", "")
        input_text = row.get("input", "")
        
        if input_text:
            question = f"{instruction}\n\nGirdi: {input_text}"
        else:
            question = instruction
        
        return question, output

    def _handle_error(self, err: Exception) -> None:
        if self.cfg.strict:
            logger.error(str(err))
            raise err
        logger.warning(str(err))

    def load_data(self):
        """
        GERİYE DÖNÜK UYUM:
        - QA_TRAIN  → List[Dict]: {"modality":"text","data":question,"target":answer}
        - TEXT_INFER→ List[Dict]: {"modality":"text","data":text}
        Not: self.load() yeni API'nin ham çıktısını döndürür; burada eski testlerin beklediği
        sözleşmeye dönüştürüyoruz.
        """
        raw = self.load()
        out = []
        if self.cfg.mode == LoadMode.QA_TRAIN:
            for item in raw:  # List[Tuple[str,str]]
                if isinstance(item, tuple) and len(item) == 2:
                    q, a = item
                    out.append({"modality": "text", "data": str(q), "target": str(a)})
        else:  # LoadMode.TEXT_INFER
            for t in raw:     # List[str]
                if isinstance(t, str) and t.strip():
                    out.append({"modality": "text", "data": t})
        return out
